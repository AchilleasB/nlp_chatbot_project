"""Text preprocessing module for document handling and chunking."""

import re
from pathlib import Path
from typing import List, Optional, Dict, Any
import docx
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from src.embeddings import CBOWModel, DocumentEmbedder
from src.vectordb.vector_db import VectorDB

from src.chunking.text_splitter import TextSplitter
from src.config import (
    DEFAULT_CHUNK_OVERLAP,
    DEFAULT_CHUNK_SIZE,
    DEFAULT_MIN_CHUNK_SIZE,
    DEFAULT_MAX_TITLE_LENGTH,
    PROCESSED_DOCUMENTS_DIR
)


class TextPreprocessor():
    """Handles document loading and text preprocessing."""
    
    def __init__(
        self,
        embedding_model: Optional[CBOWModel] = None,
        document_embedder: Optional[DocumentEmbedder] = None,
        vector_db: Optional[VectorDB] = None,
        chunk_size: int = DEFAULT_CHUNK_SIZE,
        chunk_overlap: int = DEFAULT_CHUNK_OVERLAP,
        min_chunk_size: int = DEFAULT_MIN_CHUNK_SIZE,
        max_title_length: int = DEFAULT_MAX_TITLE_LENGTH
    ):
        """Initialize the text preprocessor.
        
        Args:
            embedding_model: Optional CBOW model for word embeddings
            document_embedder: Optional document embedder for converting text to vectors
            vector_db: Optional vector database for storing embeddings
            chunk_size: Size of text chunks
            chunk_overlap: Number of characters to overlap between chunks
            min_chunk_size: Minimum size for a chunk to be considered valid
            max_title_length: Maximum length for a line to be considered a title
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.min_chunk_size = min_chunk_size
        self.max_title_length = max_title_length
        
        # Use provided models or create new ones
        self.embedding_model = embedding_model or CBOWModel()
        self.document_embedder = document_embedder or DocumentEmbedder(self.embedding_model)
        self.vector_db = vector_db or VectorDB(embedding_dim=self.document_embedder.get_embedding_dimension())
        
        self.text_splitter = TextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            min_chunk_size=min_chunk_size,
            max_title_length=max_title_length
        )
    
    def _is_cover_page(self, paragraphs: List[Paragraph]) -> bool:
        """Check if the first few paragraphs are a cover page."""
        if len(paragraphs) < 3:
            return False
        
        # Check for common cover page patterns
        first_para = paragraphs[0].text.strip()
        if not first_para:
            return False
            
        # Check for centered text (common in cover pages)
        if paragraphs[0].alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
            return True
            
        # Check for common cover page phrases
        cover_phrases = [
            r'^[A-Z][A-Z\s]{10,}$',  # ALL CAPS title
            r'^prepared by:',  # Author section
            r'^submitted to:',  # Submission info
            r'^date:',  # Date
            r'^version',  # Version number
            r'^confidential',  # Confidentiality notice
        ]
        
        return any(re.match(pattern, first_para, re.IGNORECASE) for pattern in cover_phrases)
    
    def _extract_document_metadata(self, doc: docx.Document) -> Dict[str, Any]:
        """Extract metadata from document properties."""
        core_props = doc.core_properties
        return {
            'title': core_props.title,
            'author': core_props.author,
            'created': core_props.created,
            'modified': core_props.modified,
            'last_modified_by': core_props.last_modified_by,
            'revision': core_props.revision,
            'category': core_props.category,
            'keywords': core_props.keywords,
            'comments': core_props.comments,
        }
    
    def _clean_paragraph(self, text: str) -> str:
        """Clean a single paragraph of text."""
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text)
        # Remove special characters but keep punctuation
        text = re.sub(r'[^\w\s.,!?-]', '', text)
        # Normalize whitespace
        text = ' '.join(text.split())
        return text.strip()
    
    def load_document(self, file_path: str) -> str:
        """Load a document and extract its text content.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Extracted text content
            
        Raises:
            ValueError: If file type is not supported
        """
        file_path = Path(file_path)
        
        if file_path.suffix.lower() == '.txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        elif file_path.suffix.lower() == '.docx':
            doc = docx.Document(file_path)
            
            # Skip cover page if present
            start_idx = 1 if self._is_cover_page(doc.paragraphs) else 0
            
            # Extract and clean paragraphs
            cleaned_paragraphs = []
            for para in doc.paragraphs[start_idx:]:
                text = self._clean_paragraph(para.text)
                if text:  # Only add non-empty paragraphs
                    cleaned_paragraphs.append(text)
            
            return '\n'.join(cleaned_paragraphs)
        else:
            raise ValueError(f"Unsupported file type: {file_path.suffix}")
    
    def process_text(self, text: str) -> List[str]:
        """Process text into chunks.
        
        Args:
            text: Input text
            
        Returns:
            List of text chunks
        """
        # Split text into chunks
        chunks = self.text_splitter.split_text(text)
        
        # Train model if not trained
        if not self.embedding_model.is_trained:
            print("Training embedding model on chunks...")
            self.embedding_model.train(chunks)
        
        # Generate embeddings and add to vector database
        embeddings = self.document_embedder.get_embeddings(chunks)
        self.vector_db.add_vectors(embeddings, chunks)
        
        return chunks
    
    def process_document(self, file_path: str) -> List[str]:
        """Process a document file into chunks.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            List of text chunks
        """
        text = self.load_document(file_path)
        return self.process_text(text)
    
    def process_directory(self, directory: str, file_pattern: str = "*.docx", save_processed: bool = True) -> List[str]:
        """Process all matching documents in a directory.
        
        Args:
            directory: Directory path
            file_pattern: Glob pattern for files to process
            save_processed: Whether to save processed chunks to files
            
        Returns:
            List of text chunks from all documents
        """
        directory = Path(directory)
        processed_dir = Path(PROCESSED_DOCUMENTS_DIR)
        all_chunks = []
        
        # Create processed directory if it doesn't exist
        if save_processed:
            processed_dir.mkdir(parents=True, exist_ok=True)
        
        # Process each document
        for file_path in directory.glob(file_pattern):
            try:
                print(f"\nProcessing {file_path.name}...")
                chunks = self.process_document(str(file_path))
                
                # Save processed chunks to file
                if save_processed:
                    self.save_chunks(chunks, processed_dir / f"{file_path.stem}_chunks.txt")
                
                all_chunks.extend(chunks)
                
            except Exception as e:
                print(f"Error processing {file_path}: {str(e)}")
                continue
        
        if save_processed:
            print(f"\nAll processed chunks saved to {processed_dir}")
            print(f"Total chunks across all documents: {len(all_chunks)}")
        
        return all_chunks
    
    def save_chunks(self, chunks: List[str], file_path: str) -> None:
        """Save chunks to a file."""
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(f"# Processed chunks from {file_path.name}\n")
            f.write(f"# Total chunks: {len(chunks)}\n")
            f.write("#" * 80 + "\n\n")

            for i, chunk in enumerate(chunks, 1):
                f.write(f"--- Chunk {i} ---\n")
                f.write(chunk)
                f.write("\n\n" + "#" * 40 + "\n\n")

            print(f"Saved {len(chunks)} chunks to {file_path}")
